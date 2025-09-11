import os
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional
from docx import Document

class DocumentProcessor:
    """
    Enhanced document processor for extracting content from various document formats
    """
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def extract_docx_content(self, file_path: str) -> str:
        """
        Extract text content from DOCX file using python-docx library
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            # Use python-docx for better extraction
            doc = Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            print(f"Error extracting DOCX content with python-docx: {e}")
            # Fallback to manual extraction
            return self._extract_docx_manual(file_path)
    
    def _extract_docx_manual(self, file_path: str) -> str:
        """
        Fallback manual extraction method
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Try to read the main document
                if 'word/document.xml' in docx_zip.namelist():
                    with docx_zip.open('word/document.xml') as xml_file:
                        tree = ET.parse(xml_file)
                        root = tree.getroot()
                        return self._extract_formatted_text(root)
                else:
                    return "Error: Could not find document.xml in DOCX file"
                    
        except Exception as e:
            return f"Error extracting DOCX content: {str(e)}"
    
    def _extract_formatted_text(self, root: ET.Element) -> str:
        """
        Extract text from XML with basic formatting preservation
        
        Args:
            root: Root XML element
            
        Returns:
            Extracted text with formatting
        """
        # Define namespace for Word documents
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        text_parts = []
        
        # Find all paragraphs
        paragraphs = root.findall('.//w:p', namespaces)
        
        for para in paragraphs:
            para_text = ''.join(para.itertext())
            if para_text.strip():  # Only add non-empty paragraphs
                text_parts.append(para_text.strip())
        
        # If no paragraphs found, fall back to all text using itertext
        if not text_parts:
            all_text = ''.join(root.itertext())
            if all_text.strip():
                text_parts = [all_text.strip()]
        
        # If still no content, try a different approach
        if not text_parts:
            # Try to find all text nodes recursively
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                if elem.tail and elem.tail.strip():
                    text_parts.append(elem.tail.strip())
        
        return '\n\n'.join(text_parts)
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Get metadata about the document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing metadata
        """
        try:
            file_stats = os.stat(file_path)
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_size': file_stats.st_size,
                'content_length': 0,  # Will be updated after extraction
                'extraction_success': False  # Will be updated after extraction
            }
        except Exception as e:
            return {
                'error': f"Could not get metadata: {str(e)}"
            }
    
    def process_document(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Process a document and return content with metadata
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (content, metadata)
        """
        metadata = self.get_document_metadata(file_path)
        
        if 'error' in metadata:
            return "", metadata
        
        # Extract content based on file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            content = self.extract_docx_content(file_path)
        else:
            content = f"Unsupported file format: {file_ext}"
        
        # Update metadata
        metadata['content_length'] = len(content)
        metadata['extraction_success'] = len(content) > 0 and not content.startswith('Error')
        
        return content, metadata


def extract_docx_content(file_path: str) -> str:
    """
    Convenience function to extract content from DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    processor = DocumentProcessor()
    content, _ = processor.process_document(file_path)
    return content


# Test the extraction
if __name__ == "__main__":
    file_path = "d:/NeuroLearn/NeuroLearn_V1.docx"
    
    processor = DocumentProcessor()
    content, metadata = processor.process_document(file_path)
    
    if content:
        print("Content extracted successfully!")
        print(f"Content length: {len(content)} characters")
        print("\n--- First 500 characters ---")
        print(content[:500])
        
        # Save to file
        with open("NeuroLearn_extracted_formatted.txt", "w", encoding="utf-8") as f:
            f.write(content)
        print("\nFormatted content saved to NeuroLearn_extracted_formatted.txt")
        
        print(f"\nMetadata: {metadata}")
    else:
        print("Failed to extract content")
        print(f"Metadata: {metadata}")